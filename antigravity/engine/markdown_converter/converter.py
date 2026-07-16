"""Markdown 合并与转 Word 的纯逻辑（无 UI 依赖）。

从原 ``merge_and_convert_gui.py`` 剥离，供主流水线 ``merge`` 步与独立 GUI 共用。

页序修正：OCR 产物命名为 ``{stem}_{N}.md``（如 ``微信图片_20260206125102_0.md``），
旧实现按 ``page_(\\d+)`` 排序永远匹配不到、退化为目录序。这里改为按文件名尾部
``_(\\d+)`` 数字排序，取不到时回落自然序并记 warning。
"""

from __future__ import annotations

import html
import logging
import os
import re
from pathlib import Path
from typing import List, Optional

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)

PAGE_BREAK_MARKER = "\n\n---PAGE_BREAK---\n\n"
_TRAILING_INDEX_RE = re.compile(r"_(\d+)(?:\.[A-Za-z0-9]+)?$")


def page_sort_key(filename: str):
    """排序键：支持切片 ``parent__region_N.md`` 与整页 ``stem_N.md``。

    返回 (parent, region, layout_index)；无法解析时返回 (filename, '', 0)。
    """
    name = Path(filename).name
    stem = Path(filename).stem
    layout_idx = 0
    m = _TRAILING_INDEX_RE.search(stem)
    if m:
        layout_idx = int(m.group(1))
        base = stem[: m.start()]
    else:
        base = stem
        m2 = re.search(r"page_(\d+)", stem)
        if m2:
            layout_idx = int(m2.group(1))
    if "__" in base:
        parent, region = base.split("__", 1)
        return (parent, region, layout_idx)
    return (base, "", layout_idx)


def merge_markdown_files(folder_path: str) -> Optional[str]:
    """合并文件夹内所有 markdown（跳过 *_merged.md），按页码/切片序拼接。"""
    folder = Path(folder_path)
    if not folder.is_dir():
        return None

    md_files = [f for f in os.listdir(folder_path) if f.endswith(".md") and "merged" not in f]
    if not md_files:
        return None

    md_files = sorted(md_files, key=page_sort_key)

    parts: List[str] = []
    for i, md_file in enumerate(md_files, 1):
        content = (folder / md_file).read_text(encoding="utf-8")
        # 切片页：文件名含 __ 时加可读标题，便于抽取溯源
        stem = Path(md_file).stem
        base = stem.rsplit("_", 1)[0] if re.search(r"_\d+$", stem) else stem
        if "__" in base:
            parent, region = base.split("__", 1)
            header = f"## 原图 {parent} / 区域 {region}\n\n"
            if not content.lstrip().startswith("##"):
                content = header + content
        parts.append(content)
        if i < len(md_files):
            parts.append(PAGE_BREAK_MARKER)
    return "\n".join(parts)


def clean_html_tags(content: str) -> str:
    """清理 HTML 标签，保留对齐信息标记。"""
    soup = BeautifulSoup(content, "html.parser")

    for div in soup.find_all("div"):
        text = div.get_text(strip=True)
        style = div.get("style", "")
        if "text-align: center" in style or "text-align:center" in style:
            div.replace_with(f"[CENTER]{text}[/CENTER]")
        elif "text-align: right" in style or "text-align:right" in style:
            div.replace_with(f"[RIGHT]{text}[/RIGHT]")
        else:
            div.replace_with(text)

    for img in soup.find_all("img"):
        alt_text = img.get("alt", "")
        src = img.get("src", "")
        if alt_text:
            img.replace_with(f"[图片: {alt_text}]")
        elif src:
            img.replace_with(f"[图片: {os.path.basename(src)}]")
        else:
            img.replace_with("[图片]")

    for tag in soup.find_all(["span", "p", "strong", "em", "b", "i"]):
        tag.unwrap()

    return str(soup)


_LATEX_TO_UNICODE = {
    r"\uparrow": "↑", r"\downarrow": "↓", r"\mu": "μ", r"\alpha": "α",
    r"\beta": "β", r"\gamma": "γ", r"\delta": "δ", r"\epsilon": "ε",
    r"\theta": "θ", r"\lambda": "λ", r"\sigma": "σ", r"\tau": "τ",
    r"\phi": "φ", r"\omega": "ω", r"\pm": "±", r"\times": "×",
    r"\div": "÷", r"\leq": "≤", r"\geq": "≥", r"\neq": "≠",
    r"\approx": "≈", r"\infty": "∞", r"\degree": "°", r"\cdot": "·",
    r"\%": "%", r"\$": "$", r"\{": "{", r"\}": "}", r"\_": "_",
    r"\&": "&", r"\#": "#",
}


def convert_latex_symbols(text: str) -> str:
    """把 $...$ 内的 LaTeX 符号转成 Unicode。"""
    def replace_latex(match):
        formula = match.group(1).strip()
        for latex, uni in _LATEX_TO_UNICODE.items():
            formula = formula.replace(latex, uni)
        return formula

    return re.sub(r"\$\s*(.*?)\s*\$", replace_latex, text)


_OCR_CORRECTIONS = [
    (r"征\(一\)", "征(-)"),
    (r"征\(二\)", "征(-)"),
    (r"\(一\)(?=[，。；、\s])", "(-)"),
    (r"\(二\)(?=[，。；、\s])", "(-)"),
    (r"加一", "(+)"),
    (r"减一", "(-)"),
    (r"(反射|试验|征)\(一\)", r"\1(-)"),
    (r"(Hoffmann|Babinski|Kerning|Murphy|Chvostek|Trousseau)征\(一\)", r"\1征(-)"),
]


def fix_ocr_errors(text: str) -> str:
    """修正病历里常见的 OCR 误识别（如阴性符号 (一)→(-)）。"""
    for pattern, replacement in _OCR_CORRECTIONS:
        text = re.sub(pattern, replacement, text)
    return text


def decode_html_entities(text: str) -> str:
    return html.unescape(text)


def parse_html_table(table_html: str):
    """把 HTML 表格解析为结构化数据（保留 colspan/rowspan/对齐）。"""
    soup = BeautifulSoup(table_html, "html.parser")
    table = soup.find("table")
    if not table:
        return None

    rows_data = []
    for tr in table.find_all("tr"):
        row_data = []
        for cell in tr.find_all(["td", "th"]):
            colspan = int(cell.get("colspan", 1))
            rowspan = int(cell.get("rowspan", 1))
            text = cell.get_text(strip=True)
            text = decode_html_entities(text)
            text = convert_latex_symbols(text)
            text = fix_ocr_errors(text)

            style = cell.get("style", "")
            align = "left"
            if "text-align: center" in style or "text-align:center" in style:
                align = "center"
            elif "text-align: right" in style or "text-align:right" in style:
                align = "right"

            row_data.append({
                "text": text,
                "colspan": colspan,
                "rowspan": rowspan,
                "align": align,
                "is_header": cell.name == "th",
            })
        rows_data.append(row_data)
    return rows_data


def add_table_to_doc(doc, table_data):
    """把解析后的表格写入 Word 文档，处理合并单元格。"""
    if not table_data:
        return

    max_cols = 0
    for row in table_data:
        col_count = sum(cell["colspan"] for cell in row)
        max_cols = max(max_cols, col_count)
    if max_cols == 0:
        return

    rows_count = len(table_data)
    table = doc.add_table(rows=rows_count, cols=max_cols)
    table.style = "Table Grid"

    merged_cells = set()
    for row_idx, row_data in enumerate(table_data):
        col_idx = 0
        for cell_data in row_data:
            if col_idx >= max_cols:
                break
            while (row_idx, col_idx) in merged_cells and col_idx < max_cols:
                col_idx += 1
            if col_idx >= max_cols:
                break
            try:
                cell = table.rows[row_idx].cells[col_idx]
                colspan = cell_data["colspan"]
                rowspan = cell_data["rowspan"]

                if colspan > 1 or rowspan > 1:
                    end_col = min(col_idx + colspan - 1, max_cols - 1)
                    end_row = min(row_idx + rowspan - 1, rows_count - 1)
                    for r in range(row_idx, end_row + 1):
                        for c in range(col_idx, end_col + 1):
                            if (r, c) != (row_idx, col_idx):
                                merged_cells.add((r, c))
                    if end_col > col_idx or end_row > row_idx:
                        merge_cell = table.cell(row_idx, col_idx)
                        merge_end = table.cell(end_row, end_col)
                        merge_cell.merge(merge_end)
                        cell = merge_cell

                cell.text = cell_data["text"]
                for paragraph in cell.paragraphs:
                    if cell_data["align"] == "center":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif cell_data["align"] == "right":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                col_idx += colspan
            except Exception as exc:
                logger.warning("处理表格单元格 (%d,%d) 出错: %s", row_idx, col_idx, exc)
                col_idx += 1


def _set_song_font(run_or_style):
    run_or_style.font.name = "宋体"
    run_or_style.font.size = Pt(10.5)
    run_or_style._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
    )


def convert_md_to_docx(md_content: str, output_path: str) -> bool:
    """把合并后的 Markdown（含 HTML 表格）转为 Word 文档。"""
    try:
        doc = Document()
        _set_song_font(doc.styles["Normal"])

        pages = md_content.split("---PAGE_BREAK---")
        for page_idx, page_content in enumerate(pages):
            parts = re.split(r"(<table[^>]*>.*?</table>)", page_content, flags=re.DOTALL | re.IGNORECASE)
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                if part.lower().startswith("<table"):
                    table_data = parse_html_table(part)
                    if table_data:
                        add_table_to_doc(doc, table_data)
                        doc.add_paragraph()
                else:
                    cleaned = clean_html_tags(part)
                    cleaned = decode_html_entities(cleaned)
                    cleaned = convert_latex_symbols(cleaned)
                    cleaned = fix_ocr_errors(cleaned)
                    for line in cleaned.split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        align = WD_ALIGN_PARAGRAPH.LEFT
                        if line.startswith("[CENTER]") and line.endswith("[/CENTER]"):
                            line = line[8:-9]
                            align = WD_ALIGN_PARAGRAPH.CENTER
                        elif line.startswith("[RIGHT]") and line.endswith("[/RIGHT]"):
                            line = line[7:-8]
                            align = WD_ALIGN_PARAGRAPH.RIGHT

                        if line.startswith("# "):
                            p = doc.add_paragraph(line[2:]); p.style = "Heading 1"
                        elif line.startswith("## "):
                            p = doc.add_paragraph(line[3:]); p.style = "Heading 2"
                        elif line.startswith("### "):
                            p = doc.add_paragraph(line[4:]); p.style = "Heading 3"
                        else:
                            p = doc.add_paragraph(line)
                        p.alignment = align
                        for run in p.runs:
                            _set_song_font(run)
            if page_idx < len(pages) - 1:
                doc.add_page_break()

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return True
    except Exception as exc:
        logger.error("Markdown 转 Word 失败: %s", exc, exc_info=True)
        return False


def merge_patient_folder(folder: str, make_docx: bool = True) -> Optional[Path]:
    """合并单个病人目录的逐页 md，写出 ``{名字}_merged.md``（可选 docx）。

    返回 merged.md 路径；目录内无可合并 md 时返回 None。
    """
    folder_path = Path(folder)
    merged = merge_markdown_files(str(folder_path))
    if merged is None:
        return None

    name = folder_path.name
    merged_md = folder_path / f"{name}_merged.md"
    merged_md.write_text(merged, encoding="utf-8")

    if make_docx:
        docx_path = folder_path / f"{name}.docx"
        convert_md_to_docx(merged, str(docx_path))
    return merged_md
